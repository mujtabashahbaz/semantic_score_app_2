from docx import Document
import math

def divide_docx(input_path, output_prefix, num_chunks=10):
    """
    Divide a .docx file into multiple chunks and save them as separate files.
    
    Args:
        input_path (str): Path to the input .docx file
        output_prefix (str): Prefix for output files (e.g., "output_" will create "output_1.docx", etc.)
        num_chunks (int): Number of chunks to divide into (default: 10)
    """
    # Load the document
    doc = Document(input_path)
    
    # Calculate paragraphs per chunk
    total_paragraphs = len(doc.paragraphs)
    paragraphs_per_chunk = math.ceil(total_paragraphs / num_chunks)
    
    # Create and save each chunk
    for chunk_num in range(num_chunks):
        # Create a new document for this chunk
        new_doc = Document()
        
        # Calculate start and end indices for this chunk
        start_idx = chunk_num * paragraphs_per_chunk
        end_idx = min((chunk_num + 1) * paragraphs_per_chunk, total_paragraphs)
        
        # Copy paragraphs to the new document
        for para in doc.paragraphs[start_idx:end_idx]:
            new_doc.add_paragraph(para.text, style=para.style)
        
        # Copy tables if they exist in this paragraph range
        for table in doc.tables:
            # This is a simplified approach - table handling might need more sophistication
            # depending on your specific needs
            new_table = new_doc.add_table(rows=0, cols=len(table.columns))
            for row in table.rows:
                new_row = new_table.add_row()
                for cell in row.cells:
                    new_row.cells[row.cells.index(cell)].text = cell.text
        
        # Save the chunk
        output_path = f"{output_prefix}{chunk_num + 1}.docx"
        new_doc.save(output_path)
        print(f"Saved chunk {chunk_num + 1} to {output_path}")

# Example usage
if __name__ == "__main__":
    input_file = "input.docx"  # Replace with your input file path
    output_prefix = "chunk_"   # This will create chunk_1.docx, chunk_2.docx, etc.
    divide_docx(input_file, output_prefix, 10)